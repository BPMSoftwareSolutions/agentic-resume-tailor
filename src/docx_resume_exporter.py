"""
DOCX Resume Exporter - Export resume to DOCX format using HTML conversion.

This module provides multiple methods for converting HTML resumes to DOCX documents.
"""

import subprocess
import sys
from pathlib import Path
from typing import Optional


class DOCXResumeExporter:
    """Export resume to DOCX format using HTML conversion."""
    
    def __init__(self):
        """Initialize the DOCX exporter."""
        self.available_methods = self._check_available_methods()
    
    def _check_available_methods(self) -> dict:
        """Check which export methods are available."""
        methods = {
            'python-docx': False,
            'pandoc': False,
            'weasyprint': False
        }
        
        # Check python-docx
        try:
            import docx
            methods['python-docx'] = True
        except ImportError:
            pass
        
        # Check pandoc
        try:
            result = subprocess.run(['pandoc', '--version'], 
                                  capture_output=True, 
                                  timeout=5)
            if result.returncode == 0:
                methods['pandoc'] = True
        except (subprocess.SubprocessError, FileNotFoundError):
            pass
        
        # Check weasyprint
        try:
            import weasyprint
            methods['weasyprint'] = True
        except (ImportError, OSError):
            # OSError can occur if weasyprint dependencies aren't installed
            pass
        
        return methods
    
    def export_to_docx(self, html_path: str, output_docx_path: str) -> bool:
        """
        Convert HTML resume to DOCX document.
        
        Args:
            html_path: Path to HTML file
            output_docx_path: Path for output DOCX file
            
        Returns:
            True if successful, False otherwise
        """
        print(f"\n{'='*80}")
        print("DOCX EXPORT")
        print(f"{'='*80}\n")

        print("Checking available export methods...")
        for method, available in self.available_methods.items():
            status = "[+]" if available else "[-]"
            print(f"  {status} {method}")
        print()

        # Try methods in order of preference
        if self.available_methods['pandoc']:
            print("Attempting export with pandoc...")
            if self._try_pandoc_method(html_path, output_docx_path):
                print(f"\nDOCX export successful using pandoc!\n")
                print(f"{'='*80}\n")
                return True

        if self.available_methods['python-docx']:
            print("Attempting export with python-docx...")
            if self._try_python_docx_method(html_path, output_docx_path):
                print(f"\nDOCX export successful using python-docx!\n")
                print(f"{'='*80}\n")
                return True

        if self.available_methods['weasyprint']:
            print("Attempting export with weasyprint (PDF intermediate)...")
            if self._try_pdf_conversion_method(html_path, output_docx_path):
                print(f"\nDOCX export successful using weasyprint!\n")
                print(f"{'='*80}\n")
                return True

        print("\nNo export methods available!")
        print("\nTo enable DOCX export, install one of the following:")
        print("  - pandoc: https://pandoc.org/installing.html")
        print("  - python-docx: pip install python-docx beautifulsoup4")
        print("  - weasyprint: pip install weasyprint\n")
        print(f"{'='*80}\n")
        
        return False
    
    def _try_pandoc_method(self, html_path: str, output_docx_path: str) -> bool:
        """Try using pandoc for conversion."""
        try:
            result = subprocess.run([
                'pandoc',
                html_path,
                '-o', output_docx_path,
                '--from=html',
                '--to=docx'
            ], capture_output=True, text=True, timeout=30)
            
            if result.returncode == 0:
                return True
            else:
                print(f"WARNING: Pandoc error: {result.stderr}")
                return False

        except Exception as e:
            print(f"WARNING: Pandoc method failed: {e}")
            return False
    
    def _try_python_docx_method(self, html_path: str, output_docx_path: str) -> bool:
        """
        Try using python-docx library for direct conversion.
        
        This method now properly parses hybrid HTML+SVG structure by:
        1. Extracting content from HTML elements (not SVG)
        2. Using data attributes for semantic structure
        3. Preserving headings, bullets, and formatting
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from bs4 import BeautifulSoup
            
            # Parse HTML content
            with open(html_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Create document
            doc = Document()
            
            # Set narrow margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(0.5)
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
            
            # Find resume container (hybrid HTML structure)
            container = soup.find(class_='resume-container')
            if not container:
                # Fallback: try to parse SVG (legacy support)
                return self._try_svg_text_extraction(soup, doc, output_docx_path)
            
            # Extract header section
            header = container.find(attrs={'data-section': 'header'})
            if not header:
                header = container.find(class_='header')
            
            if header:
                # Extract name
                name_elem = header.find(attrs={'data-field': 'name'})
                if not name_elem:
                    name_elem = header.find('h1')
                
                if name_elem:
                    p = doc.add_paragraph(name_elem.get_text().strip())
                    p.style = 'Heading 1'
                    run = p.runs[0]
                    run.font.size = Pt(38)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(236, 72, 153)  # Pink color
                
                # Extract title
                title_elem = header.find(attrs={'data-field': 'title'})
                if not title_elem:
                    title_elem = header.find(class_='title')
                
                if title_elem:
                    p = doc.add_paragraph(title_elem.get_text().strip())
                    run = p.runs[0]
                    run.font.size = Pt(18)
                    run.font.color.rgb = RGBColor(100, 100, 100)
                
                # Extract contact info
                contact_section = header.find(attrs={'data-section': 'contact_info'})
                if not contact_section:
                    contact_section = header.find(class_='contact-info')
                
                if contact_section:
                    contact_parts = []
                    
                    # Email
                    email_elem = contact_section.find(attrs={'data-field': 'email'})
                    if email_elem:
                        contact_parts.append(email_elem.get_text().strip())
                    
                    # Phone
                    phone_elem = contact_section.find(attrs={'data-field': 'phone'})
                    if phone_elem:
                        contact_parts.append(phone_elem.get_text().strip())
                    
                    # Location
                    location_elem = contact_section.find(attrs={'data-field': 'location'})
                    if location_elem:
                        contact_parts.append(location_elem.get_text().strip())
                    
                    if contact_parts:
                        p = doc.add_paragraph(' | '.join(contact_parts))
                        run = p.runs[0]
                        run.font.size = Pt(11)
                
                # Add spacing after header
                doc.add_paragraph()
            
            # Extract sections
            sections = container.find_all(class_='section')
            for section in sections:
                # Extract section heading
                heading_elem = section.find(class_='section-heading')
                if not heading_elem:
                    heading_elem = section.find(['h2', 'h3'])
                
                if heading_elem:
                    p = doc.add_paragraph(heading_elem.get_text().strip())
                    p.style = 'Heading 2'
                    run = p.runs[0]
                    run.font.size = Pt(18)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(236, 72, 153)  # Pink color
                
                # Extract section content based on type
                section_type = section.get('data-section', '')
                
                if section_type == 'summary' or 'summary' in section.get('class', []):
                    # Summary text
                    summary_elem = section.find(class_='summary-text')
                    if not summary_elem:
                        summary_elem = section.find('p')
                    
                    if summary_elem:
                        doc.add_paragraph(summary_elem.get_text().strip())
                
                elif section_type == 'experience' or 'experience' in section.get('class', []):
                    # Experience items
                    exp_items = section.find_all(class_='experience-item')
                    for exp_item in exp_items:
                        # Company name and dates (support both 'employer' and 'company-name')
                        company_elem = exp_item.find(class_='employer')
                        if not company_elem:
                            company_elem = exp_item.find(class_='company-name')
                        dates_elem = exp_item.find(class_='dates')
                        
                        if company_elem:
                            company_text = company_elem.get_text().strip()
                            if dates_elem:
                                company_text += f" ({dates_elem.get_text().strip()})"
                            
                            p = doc.add_paragraph(company_text)
                            run = p.runs[0]
                            run.font.size = Pt(14)
                            run.font.bold = True
                        
                        # Role and location
                        role_elem = exp_item.find(class_='role')
                        location_elem = exp_item.find(class_='location')
                        
                        if role_elem:
                            role_text = role_elem.get_text().strip()
                            if location_elem:
                                role_text += f" - {location_elem.get_text().strip()}"
                            
                            p = doc.add_paragraph(role_text)
                            run = p.runs[0]
                            run.font.size = Pt(12)
                            run.font.italic = True
                        
                        # Bullets
                        bullet_items = exp_item.find_all(class_='bullet-item')
                        for bullet_item in bullet_items:
                            bullet_text_elem = bullet_item.find(class_='bullet-text')
                            if bullet_text_elem:
                                bullet_text = bullet_text_elem.get_text().strip()
                                # Remove leading bullet if present
                                bullet_text = bullet_text.lstrip('•').strip()
                                
                                p = doc.add_paragraph(bullet_text, style='List Bullet')
                                
                                # Extract tags
                                tags_elem = bullet_item.find(class_='tags')
                                if tags_elem:
                                    tag_elements = tags_elem.find_all(class_='tag')
                                    if tag_elements:
                                        tags_text = ' | '.join([tag.get_text().strip() 
                                                               for tag in tag_elements])
                                        p = doc.add_paragraph(f"Tags: {tags_text}")
                                        run = p.runs[0]
                                        run.font.size = Pt(9)
                                        run.font.color.rgb = RGBColor(124, 58, 237)  # Purple
                        
                        # Add spacing between experience items
                        doc.add_paragraph()
                
                elif section_type == 'education' or 'education' in section.get('class', []):
                    # Education items
                    edu_items = section.find_all(class_='education-item')
                    for edu_item in edu_items:
                        # Degree (primary information)
                        degree_elem = edu_item.find(class_='degree')
                        if degree_elem:
                            p = doc.add_paragraph(degree_elem.get_text().strip())
                            run = p.runs[0]
                            run.font.size = Pt(14)
                            run.font.bold = True
                        
                        # Institution (support both 'institution' and 'school-name')
                        school_elem = edu_item.find(class_='institution')
                        if not school_elem:
                            school_elem = edu_item.find(class_='school-name')
                        
                        if school_elem:
                            p = doc.add_paragraph(school_elem.get_text().strip())
                            run = p.runs[0]
                            run.font.size = Pt(12)
                    
                    # Extract certifications subsection if present
                    certs_section = section.find(class_='certifications')
                    if certs_section:
                        # Add certifications heading
                        heading = certs_section.find(class_='subsection-heading')
                        if heading:
                            p = doc.add_paragraph(heading.get_text().strip())
                            p.style = 'Heading 3'
                            run = p.runs[0]
                            run.font.size = Pt(14)
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(236, 72, 153)
                        
                        # Add certification items
                        cert_items = certs_section.find_all(class_='certification-item')
                        for cert in cert_items:
                            doc.add_paragraph(cert.get_text().strip(), style='List Bullet')
                
                elif section_type == 'skills' or 'skills' in section.get('class', []):
                    # Skills/Core Competencies section
                    # This section can have multiple formats:
                    # 1. Two-column layout with skill categories (render as table)
                    # 2. Simple list of skill items
                    # 3. Subsections (Technical Proficiencies, Areas of Expertise)
                    
                    # Check for two-column layout
                    two_column = section.find(class_='two-column')
                    if two_column:
                        # Get both columns
                        skills_columns = two_column.find_all(class_='skills-column')
                        
                        if len(skills_columns) == 2:
                            # Create a 2-column table
                            table = doc.add_table(rows=1, cols=2)
                            table.style = 'Light Grid'  # Optional: adds subtle borders
                            
                            # Process each column
                            for col_idx, column in enumerate(skills_columns):
                                cell = table.rows[0].cells[col_idx]
                                
                                # Add subsection heading if present
                                subsection_heading = column.find(class_='subsection-heading')
                                if subsection_heading:
                                    p = cell.add_paragraph(subsection_heading.get_text().strip())
                                    run = p.runs[0]
                                    run.font.size = Pt(14)
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(124, 58, 237)  # Purple
                                
                                # Process skill categories (e.g., CLOUD: Azure, AWS...)
                                skill_categories = column.find_all(class_='skill-category')
                                for skill_cat in skill_categories:
                                    label_elem = skill_cat.find(class_='skill-label')
                                    value_elem = skill_cat.find(class_='skill-value')
                                    
                                    if label_elem and value_elem:
                                        # Create formatted paragraph in cell
                                        p = cell.add_paragraph()
                                        
                                        # Add label in bold
                                        label_run = p.add_run(label_elem.get_text().strip())
                                        label_run.font.bold = True
                                        label_run.font.size = Pt(11)
                                        
                                        # Add space
                                        p.add_run(' ')
                                        
                                        # Add value (explicitly not bold)
                                        value_run = p.add_run(value_elem.get_text().strip())
                                        value_run.font.bold = False
                                        value_run.font.size = Pt(11)
                                
                                # Process expertise list items
                                expertise_list = column.find(class_='expertise-list')
                                if expertise_list:
                                    expertise_items = expertise_list.find_all(class_='expertise-item')
                                    for item in expertise_items:
                                        p = cell.add_paragraph(item.get_text().strip(), style='List Bullet')
                                        # Ensure text is not bold
                                        for run in p.runs:
                                            run.font.bold = False
                                            run.font.size = Pt(11)
                                
                                # Remove the first empty paragraph that add_paragraph creates
                                if len(cell.paragraphs) > 0 and not cell.paragraphs[0].text.strip():
                                    cell._element.remove(cell.paragraphs[0]._element)
                        
                        else:
                            # Fallback: Process columns sequentially if not exactly 2
                            for column in skills_columns:
                                # Add subsection heading if present
                                subsection_heading = column.find(class_='subsection-heading')
                                if subsection_heading:
                                    p = doc.add_paragraph(subsection_heading.get_text().strip())
                                    p.style = 'Heading 3'
                                    run = p.runs[0]
                                    run.font.size = Pt(14)
                                    run.font.bold = True
                                    run.font.color.rgb = RGBColor(124, 58, 237)  # Purple
                                
                                # Process skill categories
                                skill_categories = column.find_all(class_='skill-category')
                                for skill_cat in skill_categories:
                                    label_elem = skill_cat.find(class_='skill-label')
                                    value_elem = skill_cat.find(class_='skill-value')
                                    
                                    if label_elem and value_elem:
                                        p = doc.add_paragraph()
                                        label_run = p.add_run(label_elem.get_text().strip())
                                        label_run.font.bold = True
                                        label_run.font.size = Pt(11)
                                        p.add_run(' ')
                                        value_run = p.add_run(value_elem.get_text().strip())
                                        value_run.font.bold = False
                                        value_run.font.size = Pt(11)
                                
                                # Process expertise list items
                                expertise_list = column.find(class_='expertise-list')
                                if expertise_list:
                                    expertise_items = expertise_list.find_all(class_='expertise-item')
                                    for item in expertise_items:
                                        p = doc.add_paragraph(item.get_text().strip(), style='List Bullet')
                                        # Ensure text is not bold
                                        for run in p.runs:
                                            run.font.bold = False
                                            run.font.size = Pt(11)
                                
                                # Add spacing between columns
                                doc.add_paragraph()
                    
                    else:
                        # Fallback: Simple skill items list
                        skill_items = section.find_all(class_='skill-item')
                        for skill_item in skill_items:
                            skill_text = skill_item.get_text().strip()
                            if skill_text:
                                doc.add_paragraph(skill_text, style='List Bullet')
                
                else:
                    # Generic section content
                    # Extract all paragraphs
                    for p_elem in section.find_all(['p', 'div'], class_=lambda c: c and 'text' in c):
                        text = p_elem.get_text().strip()
                        if text and len(text) > 3:  # Skip very short content
                            doc.add_paragraph(text)
                
                # Add spacing after section
                doc.add_paragraph()
            
            # Save document
            doc.save(output_docx_path)
            return True
            
        except ImportError:
            print("WARNING: python-docx or beautifulsoup4 not available")
            return False
        except Exception as e:
            print(f"WARNING: python-docx method failed: {e}")
            import traceback
            traceback.print_exc()
            return False
    
    def _try_svg_text_extraction(self, soup, doc, output_docx_path: str) -> bool:
        """
        Legacy fallback: Extract text from SVG elements.
        
        This is used for pure SVG resumes (not hybrid HTML+SVG).
        """
        try:
            from docx.shared import Pt
            
            svg = soup.find('svg')
            if not svg:
                return False
            
            # Find all text elements
            text_elements = svg.find_all('text')
            
            for text_elem in text_elements:
                text_content = text_elem.get_text().strip()
                if text_content:
                    # Add paragraph
                    p = doc.add_paragraph(text_content)
                    
                    # Try to apply basic styling based on font-size
                    font_size = text_elem.get('font-size', '14')
                    try:
                        size = int(font_size)
                        if size >= 24:
                            # Large text - heading
                            p.style = 'Heading 1'
                        elif size >= 16:
                            # Medium text - heading 2
                            p.style = 'Heading 2'
                    except (ValueError, TypeError):
                        pass
            
            # Save document
            doc.save(output_docx_path)
            return True

        except Exception as e:
            print(f"WARNING: SVG text extraction failed: {e}")
            return False
    
    def _try_pdf_conversion_method(self, html_path: str, output_docx_path: str) -> bool:
        """Try using weasyprint for PDF then convert to DOCX."""
        try:
            import weasyprint
            
            # Generate PDF first
            pdf_path = output_docx_path.replace('.docx', '.pdf')
            
            print(f"  -> Generating PDF: {pdf_path}")
            weasyprint.HTML(filename=html_path).write_pdf(pdf_path)

            print(f"  [+] PDF generated")
            print(f"  INFO: Note: PDF generated, but DOCX conversion requires additional tools")
            print(f"  INFO: Consider using pandoc or python-docx for direct DOCX export")
            
            # We can't easily convert PDF to DOCX without additional tools
            # Return False to try other methods
            return False
            
        except ImportError:
            print("WARNING: weasyprint not available")
            return False
        except Exception as e:
            print(f"WARNING: weasyprint method failed: {e}")
            return False


def main():
    """Main entry point for command-line usage."""
    import argparse
    
    parser = argparse.ArgumentParser(description='Export HTML resume to DOCX format')
    parser.add_argument('--input', required=True, help='Path to HTML file')
    parser.add_argument('--output', required=True, help='Path for output DOCX file')
    
    args = parser.parse_args()
    
    exporter = DOCXResumeExporter()
    success = exporter.export_to_docx(args.input, args.output)
    
    if not success:
        sys.exit(1)


if __name__ == '__main__':
    main()

