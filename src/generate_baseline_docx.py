"""
Generate baseline DOCX resume matching the original template structure.
This script creates a DOCX file with the exact structure of the original template:
- Table-based layout for headers and structured content
- Paragraph-based content for experience entries
- Specific formatting matching the original
"""
import json
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_run_font(run, size=10, bold=False, name='Calibri'):
    """Set font properties for a run."""
    run.font.name = name
    run.font.size = Pt(size)
    if bold:
        run.font.bold = True
    return run


def add_table_with_header(doc, header_text):
    """Add a 2x3 table with centered header text (yellow background)."""
    table = doc.add_table(rows=2, cols=3)
    table.autofit = False
    
    # Set column widths
    table.columns[0].width = Inches(0.5)
    table.columns[1].width = Inches(5.5)
    table.columns[2].width = Inches(0.5)
    
    # Add header text to middle column of both rows
    for row_idx in [0, 1]:
        cell = table.rows[row_idx].cells[1]
        cell.text = header_text
        
        # Center align
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_run_font(run, size=11, bold=True)
        
        # Add yellow background
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'FFFF00')  # Yellow
        cell._element.get_or_add_tcPr().append(shading_elm)
    
    return table


def generate_baseline_docx(resume_json_path, output_docx_path):
    """Generate a DOCX resume matching the original template structure."""
    
    # Load resume data
    with open(resume_json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    
    # ===== TABLE 0: Header with name and contact =====
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(3.0)
    header_table.columns[1].width = Inches(3.5)
    
    # Left cell: Name
    name_cell = header_table.rows[0].cells[0]
    name_para = name_cell.paragraphs[0]
    name_run = name_para.add_run(data['name'])
    set_run_font(name_run, size=18, bold=True)

    # Right cell: Contact info
    contact_cell = header_table.rows[0].cells[1]
    contact_para = contact_cell.paragraphs[0]
    contact_text = f"{data['contact']['email']} • {data['contact']['phone']}\nLinkedIn • {data['location']}"
    contact_run = contact_para.add_run(contact_text)
    set_run_font(contact_run, size=10)

    # ===== Title and Summary (as paragraphs) =====
    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(data['title'])
    set_run_font(title_run, size=11, bold=True)

    # Summary - first sentence bold
    summary_text = data['summary']
    sentences = summary_text.split('. ', 1)

    summary_para1 = doc.add_paragraph()
    first_sent_run = summary_para1.add_run(sentences[0] + '.')
    set_run_font(first_sent_run, size=10, bold=True)

    if len(sentences) > 1:
        summary_para2 = doc.add_paragraph()
        rest_run = summary_para2.add_run(sentences[1])
        set_run_font(rest_run, size=10)
    
    # ===== TABLE 1 & 2: Technical Proficiencies =====
    add_table_with_header(doc, "Technical Proficiencies")
    
    # Technical proficiencies content table (8 rows x 2 columns)
    tech_table = doc.add_table(rows=8, cols=2)
    tech_table.autofit = False
    tech_table.columns[0].width = Inches(2.5)
    tech_table.columns[1].width = Inches(4.0)
    
    # Add technical proficiencies data
    tech_prof = data.get('technical_proficiencies', {})
    tech_items = [
        ("Cloud Architecture & Engineering:", tech_prof.get('cloud', '')),
        ("AI & Automation:", tech_prof.get('ai', '')),
        ("DevOps & CI/CD:", tech_prof.get('devops', '')),
        ("Security & Compliance:", tech_prof.get('security', '')),
        ("Programming Languages:", tech_prof.get('languages', '')),
        ("Databases:", tech_prof.get('databases', '')),
        ("Operating Systems:", tech_prof.get('os', '')),
        ("Open Source & Innovation:", tech_prof.get('opensource', ''))
    ]
    
    for i, (label, value) in enumerate(tech_items):
        row = tech_table.rows[i]
        
        # Label cell (bold)
        label_cell = row.cells[0]
        label_para = label_cell.paragraphs[0]
        label_run = label_para.add_run(label)
        set_run_font(label_run, size=10, bold=True)

        # Value cell
        value_cell = row.cells[1]
        value_para = value_cell.paragraphs[0]
        value_run = value_para.add_run(value)
        set_run_font(value_run, size=10)
    
    # ===== TABLE 3 & 4: Areas of Expertise =====
    add_table_with_header(doc, "Areas of Expertise")
    
    # Areas of expertise content table (1 row x 3 columns)
    areas_table = doc.add_table(rows=1, cols=3)
    areas_table.autofit = False
    for col in areas_table.columns:
        col.width = Inches(2.17)
    
    # Add areas of expertise data (split into 3 columns)
    areas = data.get('areas_of_expertise', [])
    col_size = (len(areas) + 2) // 3  # Divide into 3 columns
    
    for col_idx in range(3):
        cell = areas_table.rows[0].cells[col_idx]
        start_idx = col_idx * col_size
        end_idx = min(start_idx + col_size, len(areas))
        
        cell_text = '\n'.join(areas[start_idx:end_idx])
        para = cell.paragraphs[0]
        run = para.add_run(cell_text)
        set_run_font(run, size=10)
    
    # ===== TABLE 5: Career Experience Header =====
    add_table_with_header(doc, "Career Experience")
    
    # ===== Experience entries (as paragraphs) =====
    for job in data.get('experience', []):
        # Company and dates (bold)
        company_para = doc.add_paragraph()
        company_text = f"{job['employer']}, {job.get('location', '')}      {job['dates']}"
        company_run = company_para.add_run(company_text)
        set_run_font(company_run, size=10, bold=True)

        # Role (bold)
        role_para = doc.add_paragraph()
        role_run = role_para.add_run(job['role'])
        set_run_font(role_run, size=10, bold=True)
        
        # Bullets
        for bullet in job.get('bullets', []):
            bullet_text = bullet.get('text', bullet) if isinstance(bullet, dict) else bullet
            
            # Check if bullet starts with bold project name (contains colon)
            if ':' in bullet_text:
                parts = bullet_text.split(':', 1)
                bullet_para = doc.add_paragraph()

                # Bold part (project name)
                bold_run = bullet_para.add_run(parts[0] + ':')
                set_run_font(bold_run, size=10, bold=True)

                # Normal part (description)
                normal_run = bullet_para.add_run(parts[1])
                set_run_font(normal_run, size=10)
            else:
                # Regular bullet
                bullet_para = doc.add_paragraph()
                bullet_run = bullet_para.add_run(bullet_text)
                set_run_font(bullet_run, size=10)
    
    # ===== TABLE 6: Education Header =====
    add_table_with_header(doc, "Education & Professional Development")
    
    # ===== Education entries (as paragraphs) =====
    for edu in data.get('education', []):
        edu_para = doc.add_paragraph()
        edu_text = f"{edu.get('degree', '')} | {edu.get('institution', '')}, {edu.get('location', '')}"
        edu_run = edu_para.add_run(edu_text)
        set_run_font(edu_run, size=10)

    # ===== Certifications (as paragraphs) =====
    for cert in data.get('certifications', []):
        cert_para = doc.add_paragraph()
        cert_run = cert_para.add_run(cert)
        set_run_font(cert_run, size=10)

    # ===== Achievements (as paragraphs) =====
    for achievement in data.get('achievements', []):
        ach_para = doc.add_paragraph()
        ach_run = ach_para.add_run(achievement)
        set_run_font(ach_run, size=10)
    
    # Save document
    doc.save(output_docx_path)
    print(f"✅ Generated baseline DOCX: {output_docx_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate baseline DOCX resume")
    parser.add_argument("--json", required=True, help="Path to master resume JSON file")
    parser.add_argument("--docx", required=True, help="Path to output DOCX file")
    args = parser.parse_args()
    
    generate_baseline_docx(args.json, args.docx)

