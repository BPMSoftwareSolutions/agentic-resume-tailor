import argparse, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def md_to_docx(md_path, docx_path):
    """Convert Markdown resume to ATS-friendly DOCX format matching the template."""
    md_text = Path(md_path).read_text(encoding='utf-8')
    doc = Document()

    # Set default font for ATS compatibility
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)  # Default 10pt

    lines = md_text.split('\n')
    i = 0
    current_section = None
    summary_first_line = True
    seen_title = False

    while i < len(lines):
        line = lines[i].strip()

        # Skip empty lines
        if not line:
            i += 1
            continue

        # H1 - Name (skip, we'll use title line instead)
        if line.startswith('# '):
            i += 1
            continue

        # H2 - Section headers (track section, but don't output)
        elif line.startswith('## '):
            current_section = line[3:].lower()
            if current_section == 'summary':
                summary_first_line = True
            i += 1
            continue

        # Title line (DevOps & Automation Engineer...) - first non-header line
        elif not seen_title and current_section is None and not line.startswith('#'):
            # This is the title line
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.size = Pt(13)
            run.font.bold = True
            seen_title = True

        # Contact info line (second non-header line before any section)
        elif seen_title and current_section is None and not line.startswith('#'):
            # This is contact info - not bold
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.size = Pt(10)

        # Summary section - split into sentences, first bold
        elif current_section == 'summary':
            # Split long summary into multiple sentences
            sentences = re.split(r'(?<=[.!?])\s+', line)
            for j, sentence in enumerate(sentences):
                if sentence.strip():
                    p = doc.add_paragraph(sentence.strip())
                    run = p.runs[0]
                    run.font.size = Pt(10)
                    if j == 0 and summary_first_line:
                        run.font.bold = True
                        summary_first_line = False

        # Skills section - skip bullets
        elif current_section == 'skills' and line.startswith('- '):
            i += 1
            continue

        # Experience section
        elif current_section == 'experience':
            # Job title line with dates: **Role — Company** (Dates)
            if line.startswith('**') and '**' in line[2:]:
                # Parse: **Role — Company** (Dates)
                match = re.match(r'\*\*(.+?)\s*—\s*(.+?)\*\*\s*\((.+?)\)', line)
                if match:
                    role, company, dates = match.groups()
                    # Company and dates line (bold)
                    p = doc.add_paragraph(f"{company}      {dates}")
                    run = p.runs[0]
                    run.font.bold = True
                    run.font.size = Pt(10)

                    # Role line (bold)
                    p = doc.add_paragraph(role)
                    run = p.runs[0]
                    run.font.bold = True
                    run.font.size = Pt(10)
                else:
                    # Fallback: just make it bold
                    text = line.replace('**', '')
                    p = doc.add_paragraph(text)
                    run = p.runs[0]
                    run.font.bold = True
                    run.font.size = Pt(10)

            # Bullet points - convert to regular paragraphs (no bullets)
            elif line.startswith('- '):
                text = line[2:]

                # Check if bullet starts with a project name (contains colon)
                if ':' in text:
                    # Bold the project name part
                    parts = text.split(':', 1)
                    p = doc.add_paragraph()
                    run = p.add_run(parts[0] + ':')
                    run.font.bold = True
                    run.font.size = Pt(10)
                    if len(parts) > 1:
                        run = p.add_run(parts[1])
                        run.font.size = Pt(10)
                else:
                    # Regular bullet text
                    p = doc.add_paragraph(text)
                    run = p.runs[0]
                    run.font.size = Pt(10)

        # Education section - make bold
        elif current_section == 'education':
            if line.startswith('- '):
                text = line[2:]
                p = doc.add_paragraph(text)
                run = p.runs[0]
                run.font.bold = True
                run.font.size = Pt(10)
            else:
                p = doc.add_paragraph(line)
                run = p.runs[0]
                run.font.bold = True
                run.font.size = Pt(10)

        # Regular text (contact info, etc.)
        else:
            p = doc.add_paragraph(line)
            if p.runs:
                run = p.runs[0]
                run.font.size = Pt(10)

        i += 1

    doc.save(docx_path)
    print(f"Saved ATS-friendly DOCX: {docx_path}")

if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--md", required=True)
    ap.add_argument("--docx", required=True)
    args = ap.parse_args()
    md_to_docx(args.md, args.docx)
