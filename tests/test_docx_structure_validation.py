"""
Comprehensive DOCX structure validation tests.
These tests validate that generated DOCX files match the structure of the original template.
"""
import sys
from pathlib import Path
import tempfile
import os

# Add src to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / 'src'))

import unittest
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class TestDocxStructureValidation(unittest.TestCase):
    """Test that generated DOCX matches the original template structure."""
    
    def setUp(self):
        """Load the generated output and reference template."""
        # Use the generated baseline for testing
        self.output_path = Path(__file__).parent.parent / "out" / "test_generated_baseline.docx"
        # Use the original template as reference
        self.reference_path = Path(__file__).parent.parent / "out" / "Sidney Jones Resume - Solution Architect Leader.docx"

        if not self.output_path.exists():
            self.skipTest(f"Output file not found: {self.output_path}")

        if not self.reference_path.exists():
            self.skipTest(f"Reference file not found: {self.reference_path}")

        self.output_doc = Document(str(self.output_path))
        self.reference_doc = Document(str(self.reference_path))
    
    def test_document_has_expected_sections(self):
        """Document should have at least one section."""
        out_sections = len(self.output_doc.sections)

        self.assertGreaterEqual(
            out_sections,
            1,
            f"Document should have at least 1 section, found {out_sections}"
        )

    def test_has_expected_number_of_tables(self):
        """Document should have 7 tables matching the original structure."""
        out_tables = len(self.output_doc.tables)

        self.assertEqual(
            7,
            out_tables,
            f"Expected 7 tables (header, tech prof header, tech prof content, areas header, areas content, career header, education header), found {out_tables}"
        )

    def test_table_0_is_header_with_name_and_contact(self):
        """Table 0 should be 1x2 header table with name and contact info."""
        if len(self.output_doc.tables) < 1:
            self.fail("Document has no tables")

        header_table = self.output_doc.tables[0]

        self.assertEqual(
            len(header_table.rows),
            1,
            f"Header table should have 1 row, found {len(header_table.rows)}"
        )

        self.assertEqual(
            len(header_table.columns),
            2,
            f"Header table should have 2 columns, found {len(header_table.columns)}"
        )

    def test_table_1_is_technical_proficiencies_header(self):
        """Table 1 should be 2x3 Technical Proficiencies header."""
        if len(self.output_doc.tables) < 2:
            self.fail("Document has less than 2 tables")

        table = self.output_doc.tables[1]

        self.assertEqual(
            len(table.rows),
            2,
            f"Technical Proficiencies header should have 2 rows, found {len(table.rows)}"
        )

        self.assertEqual(
            len(table.columns),
            3,
            f"Technical Proficiencies header should have 3 columns, found {len(table.columns)}"
        )

    def test_table_2_is_technical_proficiencies_content(self):
        """Table 2 should be 8x2 Technical Proficiencies content."""
        if len(self.output_doc.tables) < 3:
            self.fail("Document has less than 3 tables")

        table = self.output_doc.tables[2]

        self.assertEqual(
            len(table.rows),
            8,
            f"Technical Proficiencies content should have 8 rows, found {len(table.rows)}"
        )

        self.assertEqual(
            len(table.columns),
            2,
            f"Technical Proficiencies content should have 2 columns, found {len(table.columns)}"
        )
    
    def test_header_table_contains_name(self):
        """First table (header) should contain the name 'Sidney Jones'."""
        if len(self.output_doc.tables) == 0:
            self.fail("Output document has no tables")
        
        header_table = self.output_doc.tables[0]
        header_text = ""
        for row in header_table.rows:
            for cell in row.cells:
                header_text += cell.text + " "
        
        self.assertIn(
            "Sidney Jones",
            header_text,
            f"Name 'Sidney Jones' not found in header table. Found: {header_text[:200]}"
        )
    
    def test_header_table_contains_contact_info(self):
        """First table (header) should contain contact information."""
        if len(self.output_doc.tables) == 0:
            self.fail("Output document has no tables")

        header_table = self.output_doc.tables[0]
        header_text = ""
        for row in header_table.rows:
            for cell in row.cells:
                header_text += cell.text + " "

        # Check for email (should contain an @ symbol and domain)
        self.assertTrue(
            "@" in header_text and (".com" in header_text or ".org" in header_text),
            f"Email not found in header table. Found: {header_text[:200]}"
        )

        # Check for phone (should contain digits)
        self.assertTrue(
            "248" in header_text and "802" in header_text,
            f"Phone not found in header table. Found: {header_text[:200]}"
        )
    
    def test_technical_proficiencies_section_exists(self):
        """Document should have a 'Technical Proficiencies' section."""
        all_text = ""
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        
        self.assertIn(
            "Technical Proficiencies",
            all_text,
            "Technical Proficiencies section not found"
        )
    
    def test_areas_of_expertise_section_exists(self):
        """Document should have an 'Areas of Expertise' section."""
        all_text = ""
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        
        self.assertIn(
            "Areas of Expertise",
            all_text,
            "Areas of Expertise section not found"
        )
    
    def test_career_experience_section_exists(self):
        """Document should have a 'Career Experience' section."""
        all_text = ""
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        
        self.assertIn(
            "Career Experience",
            all_text,
            "Career Experience section not found"
        )
    
    def test_education_section_exists(self):
        """Document should have an 'Education & Professional Development' section."""
        all_text = ""
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "
        
        self.assertIn(
            "Education",
            all_text,
            "Education section not found"
        )
    
    def test_job_titles_present(self):
        """Document should contain key job titles from the resume."""
        all_text = ""
        # Check both tables and paragraphs
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text + " "

        for para in self.output_doc.paragraphs:
            all_text += para.text + " "

        # Check for at least one job title
        job_titles = [
            "BPM Software Solutions",
            "Interactive Business Solutions",
            "Soare Enterprises",
            "John Deere Landscapes"
        ]

        found_titles = [title for title in job_titles if title in all_text]

        self.assertGreater(
            len(found_titles),
            0,
            f"No job titles found in document. Expected one of: {job_titles}. Found text: {all_text[:500]}"
        )
    
    def test_visual_formatting_preserved(self):
        """Check that visual formatting elements are preserved."""
        # This is a basic check - we verify that the document has styled content
        has_bold = False
        has_colored_text = False

        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.bold:
                                has_bold = True
                            if run.font.color and run.font.color.rgb:
                                has_colored_text = True

        self.assertTrue(
            has_bold,
            "No bold text found - formatting may not be preserved"
        )

    def test_technical_proficiencies_data_populated(self):
        """Technical proficiencies table should have actual data in the value column."""
        if len(self.output_doc.tables) < 3:
            self.fail("Document doesn't have enough tables for technical proficiencies")

        tech_table = self.output_doc.tables[2]  # Table 2 is tech prof content

        # Check that at least 4 rows have non-empty content in the VALUE column (column 1)
        rows_with_content = 0
        for row in tech_table.rows:
            if len(row.cells) >= 2:
                value_cell = row.cells[1]  # Second column should have the values
                if value_cell.text.strip() and len(value_cell.text.strip()) > 5:
                    rows_with_content += 1

        self.assertGreaterEqual(
            rows_with_content,
            4,
            f"Technical proficiencies table should have at least 4 rows with content, found {rows_with_content}"
        )

    def test_areas_of_expertise_data_populated(self):
        """Areas of expertise table should have actual data, not empty cells."""
        if len(self.output_doc.tables) < 5:
            self.fail("Document doesn't have enough tables for areas of expertise")

        areas_table = self.output_doc.tables[4]  # Table 4 is areas content

        # Check that at least one cell has non-empty content
        has_content = False
        for row in areas_table.rows:
            for cell in row.cells:
                if cell.text.strip() and len(cell.text.strip()) > 10:
                    has_content = True
                    break

        self.assertTrue(
            has_content,
            "Areas of expertise table appears to be empty or has minimal content"
        )

    def test_experience_bullets_present(self):
        """Document should have experience bullet points."""
        # Count paragraphs with experience content
        experience_paras = 0
        found_experience_section = False

        for para in self.output_doc.paragraphs:
            text = para.text.strip()
            # Look for company names or role indicators
            if any(company in text for company in ["BPM Software Solutions", "Interactive Business Solutions", "Soare Enterprises"]):
                found_experience_section = True
            # Count non-empty paragraphs after finding experience section
            if found_experience_section and text and len(text) > 20:
                experience_paras += 1

        self.assertGreater(
            experience_paras,
            5,
            f"Expected at least 5 experience paragraphs, found {experience_paras}"
        )
    
    def test_document_not_empty(self):
        """Output document should not be empty."""
        total_text = ""
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_text += cell.text
        
        for para in self.output_doc.paragraphs:
            total_text += para.text
        
        self.assertGreater(
            len(total_text.strip()),
            100,
            f"Document appears to be empty or too short. Length: {len(total_text)}"
        )


if __name__ == '__main__':
    unittest.main()

