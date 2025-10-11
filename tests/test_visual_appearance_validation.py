"""
Visual appearance validation tests.
These tests specifically target the visual rendering issues that make
documents look unprofessional despite having correct content and structure.
"""
import sys
from pathlib import Path
import tempfile
import os

# Add src to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / 'src'))

import unittest
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class TestVisualAppearanceValidation(unittest.TestCase):
    """Test that documents will render with professional visual appearance."""
    
    def setUp(self):
        """Load the generated output and reference template."""
        # Use the generated baseline for testing
        self.output_path = Path(__file__).parent.parent / "out" / "test_generated_baseline.docx"
        # Use the original template as reference
        self.reference_path = Path(__file__).parent.parent / "out" / "Sidney Jones Resume - Solution Architect Leader.docx"

        if not self.output_path.exists():
            self.skipTest(f"Output file not found: {self.output_path}")

        if not self.reference_path.exists():
            self.skipTest(f"Reference file not found: {self.reference_path}")

        self.output_doc = Document(str(self.output_path))
        self.reference_doc = Document(str(self.reference_path))
    
    def test_table_cell_padding_and_spacing(self):
        """Tables should have proper cell padding to avoid cramped appearance."""
        appearance_issues = []
        
        for table_idx, table in enumerate(self.output_doc.tables):
            # Check if tables have reasonable structure for visual appeal
            if len(table.rows) == 0 or len(table.columns) == 0:
                appearance_issues.append(f"Table {table_idx}: Empty table structure")
                continue
            
            # Check for overly dense content in cells
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        # Very long text in a single cell without breaks
                        if len(cell_text) > 500 and '\n' not in cell_text:
                            appearance_issues.append(
                                f"Table {table_idx}, Cell ({row_idx},{cell_idx}): "
                                f"Very long text without breaks ({len(cell_text)} chars)"
                            )
                        
                        # Check for potential formatting that could look cramped
                        if len(cell.paragraphs) == 1 and len(cell_text) > 200:
                            # Single paragraph with lots of content might look cramped
                            para = cell.paragraphs[0]
                            
                            # Check if paragraph has any spacing formatting
                            has_spacing = (
                                para.paragraph_format.space_before or
                                para.paragraph_format.space_after or
                                para.paragraph_format.line_spacing
                            )
                            
                            if not has_spacing:
                                appearance_issues.append(
                                    f"Table {table_idx}, Cell ({row_idx},{cell_idx}): "
                                    f"Long content without spacing formatting"
                                )
        
        if appearance_issues:
            # Limit to most critical issues
            if len(appearance_issues) > 5:
                appearance_issues = appearance_issues[:5] + [f"... and {len(appearance_issues) - 5} more issues"]
            self.fail(f"Table visual appearance issues:\n" + "\n".join(appearance_issues))
    
    def test_font_rendering_consistency(self):
        """Fonts should be set explicitly to ensure consistent rendering."""
        font_issues = []
        
        # Check for undefined or inconsistent font settings
        undefined_fonts = 0
        font_families = set()
        
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.name is None:
                                undefined_fonts += 1
                            elif run.font.name:
                                font_families.add(run.font.name)
        
        for paragraph in self.output_doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name is None:
                    undefined_fonts += 1
                elif run.font.name:
                    font_families.add(run.font.name)
        
        if undefined_fonts > 10:  # Some undefined fonts are normal, but too many indicate issues
            font_issues.append(f"Many undefined fonts ({undefined_fonts}) - may cause inconsistent rendering")
        
        # Check for web-safe professional fonts
        professional_fonts = {
            'Arial', 'Calibri', 'Times New Roman', 'Georgia', 
            'Verdana', 'Helvetica', 'Tahoma', 'Trebuchet MS'
        }
        
        non_professional_fonts = font_families - professional_fonts
        if len(non_professional_fonts) > 0:
            font_issues.append(f"Non-standard fonts used: {list(non_professional_fonts)} - may not render consistently")
        
        if font_issues:
            self.fail(f"Font rendering issues:\n" + "\n".join(font_issues))
    
    def test_color_and_styling_consistency(self):
        """Colors and styling should be applied consistently for professional appearance."""
        styling_issues = []
        
        # Collect color information
        colors_used = set()
        inconsistent_styling = []
        
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            # Check font colors
                            if run.font.color and run.font.color.rgb:
                                colors_used.add(str(run.font.color.rgb))
                            
                            # Check for inconsistent styling patterns
                            if run.bold and run.italic and run.underline:
                                inconsistent_styling.append("Text with bold+italic+underline (may be over-formatted)")
        
        # Too many colors can look unprofessional
        if len(colors_used) > 4:
            styling_issues.append(f"Too many text colors used ({len(colors_used)}) - may look unprofessional")
        
        # Check for styling consistency
        if len(inconsistent_styling) > 3:
            styling_issues.append(f"Over-formatted text detected in {len(inconsistent_styling)} places")
        
        if styling_issues:
            self.fail(f"Color and styling issues:\n" + "\n".join(styling_issues))
    
    def test_table_border_and_structure_quality(self):
        """Tables should have proper borders and structure for professional appearance."""
        table_issues = []
        
        for table_idx, table in enumerate(self.output_doc.tables):
            # Check table structure that affects visual appearance
            if len(table.rows) > 0 and len(table.columns) > 0:
                
                # Check for tables with too many empty cells (looks sparse)
                total_cells = len(table.rows) * len(table.columns)
                empty_cells = 0
                
                for row in table.rows:
                    for cell in row.cells:
                        if not cell.text.strip():
                            empty_cells += 1
                
                empty_ratio = empty_cells / total_cells
                # Allow header tables (2x3 with centered content) to have empty spacer cells
                is_header_table = (len(table.rows) == 2 and len(table.columns) == 3)
                if empty_ratio > 0.6 and not is_header_table:  # More than 60% empty
                    table_issues.append(f"Table {table_idx}: Mostly empty ({empty_ratio:.1%}) - may look sparse")
                
                # Check for tables with uneven content distribution
                cell_lengths = []
                for row in table.rows:
                    for cell in row.cells:
                        cell_lengths.append(len(cell.text.strip()))
                
                if cell_lengths:
                    max_length = max(cell_lengths)
                    non_empty_lengths = [l for l in cell_lengths if l > 0]

                    if non_empty_lengths and max_length > 0:
                        # Very uneven content distribution can look bad
                        # But allow label/value tables (like tech proficiencies) to have uneven distribution
                        is_label_value_table = (len(table.columns) == 2 and len(table.rows) >= 4)
                        if max_length / min(non_empty_lengths) > 10 and not is_label_value_table:
                            table_issues.append(f"Table {table_idx}: Very uneven content distribution")
        
        if table_issues:
            self.fail(f"Table structure appearance issues:\n" + "\n".join(table_issues))
    
    def test_prevents_document_corruption_indicators(self):
        """Check for indicators that the document might be corrupted or malformed."""
        corruption_indicators = []
        
        # Check for extremely unusual document structure
        total_content = 0
        total_tables = len(self.output_doc.tables)
        total_paragraphs = len(self.output_doc.paragraphs)
        
        # Count actual content
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total_content += len(cell.text.strip())
        
        for paragraph in self.output_doc.paragraphs:
            total_content += len(paragraph.text.strip())
        
        # Unusual ratios might indicate problems
        if total_tables > 0 and total_content > 0:
            content_per_table = total_content / total_tables
            
            if content_per_table < 50:  # Very little content per table
                corruption_indicators.append(f"Very little content per table ({content_per_table:.1f} chars)")
            
            if content_per_table > 2000:  # Too much content per table
                corruption_indicators.append(f"Excessive content per table ({content_per_table:.1f} chars)")
        
        # Check for duplicate content (might indicate generation errors)
        all_cell_texts = []
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and len(text) > 20:  # Only check substantial content
                        all_cell_texts.append(text)
        
        if len(all_cell_texts) > 0:
            unique_texts = set(all_cell_texts)
            duplicate_ratio = 1 - (len(unique_texts) / len(all_cell_texts))
            
            if duplicate_ratio > 0.3:  # More than 30% duplicates
                corruption_indicators.append(f"High content duplication ({duplicate_ratio:.1%}) - may indicate generation errors")
        
        if corruption_indicators:
            self.fail(f"Document corruption indicators:\n" + "\n".join(corruption_indicators))
    
    def test_compare_visual_complexity_with_reference(self):
        """Generated document should have similar visual complexity to reference."""
        output_complexity = self._calculate_visual_complexity(self.output_doc)
        reference_complexity = self._calculate_visual_complexity(self.reference_doc)
        
        complexity_issues = []
        
        # Compare complexity metrics
        complexity_diff = abs(output_complexity['total_score'] - reference_complexity['total_score'])
        complexity_ratio = complexity_diff / reference_complexity['total_score']
        
        if complexity_ratio > 0.5:  # More than 50% difference
            complexity_issues.append(
                f"Visual complexity very different from reference: "
                f"output={output_complexity['total_score']:.1f}, "
                f"reference={reference_complexity['total_score']:.1f}"
            )
        
        # Check specific complexity aspects
        if output_complexity['table_complexity'] < reference_complexity['table_complexity'] * 0.7:
            complexity_issues.append("Tables significantly less complex than reference")
        
        if output_complexity['formatting_complexity'] < reference_complexity['formatting_complexity'] * 0.7:
            complexity_issues.append("Formatting significantly less complex than reference")
        
        if complexity_issues:
            self.fail(f"Visual complexity issues:\n" + "\n".join(complexity_issues))
    
    def _calculate_visual_complexity(self, doc):
        """Calculate visual complexity metrics for a document."""
        complexity = {
            'table_count': len(doc.tables),
            'table_complexity': 0,
            'formatting_complexity': 0,
            'content_variety': 0,
            'total_score': 0
        }
        
        # Table complexity
        for table in doc.tables:
            table_size = len(table.rows) * len(table.columns)
            complexity['table_complexity'] += table_size
        
        # Formatting complexity
        font_families = set()
        font_sizes = set()
        has_bold = False
        has_colors = False
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.name:
                                font_families.add(run.font.name)
                            if run.font.size:
                                font_sizes.add(run.font.size.pt)
                            if run.bold:
                                has_bold = True
                            if run.font.color and run.font.color.rgb:
                                has_colors = True
        
        complexity['formatting_complexity'] = (
            len(font_families) * 2 +
            len(font_sizes) * 1.5 +
            (10 if has_bold else 0) +
            (5 if has_colors else 0)
        )
        
        # Content variety (different types of content)
        all_text = ""
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text.lower() + " "
        
        content_indicators = ['experience', 'education', 'skills', 'technical', 'proficiencies', 'summary']
        complexity['content_variety'] = sum(1 for indicator in content_indicators if indicator in all_text) * 3
        
        # Total complexity score
        complexity['total_score'] = (
            complexity['table_complexity'] +
            complexity['formatting_complexity'] +
            complexity['content_variety']
        )
        
        return complexity


if __name__ == '__main__':
    unittest.main()