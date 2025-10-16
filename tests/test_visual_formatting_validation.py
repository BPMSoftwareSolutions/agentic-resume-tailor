"""
Comprehensive visual formatting validation tests.
These tests validate that generated DOCX files have proper professional formatting
that matches the visual quality of the reference document.
"""

import os
import sys
import tempfile
from pathlib import Path

# Add src to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

import unittest

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


class TestVisualFormattingValidation(unittest.TestCase):
    """Test that generated DOCX has professional visual formatting."""

    def setUp(self):
        """Load the generated output and reference template."""
        # Use the generated baseline for testing
        self.output_path = (
            Path(__file__).parent.parent / "out" / "test_generated_baseline.docx"
        )
        # Use the original template as reference
        self.reference_path = (
            Path(__file__).parent.parent
            / "out"
            / "Sidney Jones Resume - Solution Architect Leader.docx"
        )

        if not self.output_path.exists():
            self.skipTest(f"Output file not found: {self.output_path}")

        if not self.reference_path.exists():
            self.skipTest(f"Reference file not found: {self.reference_path}")

        self.output_doc = Document(str(self.output_path))
        self.reference_doc = Document(str(self.reference_path))

    def test_font_sizes_are_consistent(self):
        """Font sizes should be consistent and professional throughout the document."""
        font_sizes = []

        # Collect font sizes from output document
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.size:
                                font_sizes.append(run.font.size.pt)

        for paragraph in self.output_doc.paragraphs:
            for run in paragraph.runs:
                if run.font.size:
                    font_sizes.append(run.font.size.pt)

        # Check that we have reasonable font sizes (typically 10-16pt for professional docs)
        if font_sizes:
            min_size = min(font_sizes)
            max_size = max(font_sizes)

            self.assertGreaterEqual(
                min_size,
                8,
                f"Font sizes too small (min: {min_size}pt). Professional documents should use 8pt+ fonts.",
            )

            self.assertLessEqual(
                max_size,
                20,
                f"Font sizes too large (max: {max_size}pt). Professional documents should use fonts under 20pt.",
            )

            # Should have limited variety in font sizes (typically 2-4 different sizes)
            unique_sizes = len(set(font_sizes))
            self.assertLessEqual(
                unique_sizes,
                6,
                f"Too many different font sizes ({unique_sizes}). Professional documents use consistent sizing.",
            )

    def test_table_cell_alignment_consistency(self):
        """Table cells should have consistent and appropriate alignment."""
        alignment_issues = []

        for table_idx, table in enumerate(self.output_doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, paragraph in enumerate(cell.paragraphs):
                        # Check for mixed alignment within tables
                        if paragraph.alignment is None:
                            # Default alignment is fine, but track it
                            continue

                        # Document any unusual alignments for review
                        if paragraph.alignment not in [
                            WD_ALIGN_PARAGRAPH.LEFT,
                            WD_ALIGN_PARAGRAPH.CENTER,
                            None,
                        ]:
                            alignment_issues.append(
                                f"Table {table_idx}, Row {row_idx}, Cell {cell_idx}, Para {para_idx}: "
                                f"Unusual alignment {paragraph.alignment}"
                            )

        # This is more of a warning than a failure - unusual alignments may be intentional
        if len(alignment_issues) > 5:
            self.fail(
                f"Many alignment inconsistencies found:\n"
                + "\n".join(alignment_issues[:10])
            )

    def test_table_structure_consistency(self):
        """Tables should have consistent structure and proper formatting."""
        table_issues = []

        for table_idx, table in enumerate(self.output_doc.tables):
            # Check for empty or malformed tables
            if len(table.rows) == 0:
                table_issues.append(f"Table {table_idx}: No rows")
                continue

            if len(table.columns) == 0:
                table_issues.append(f"Table {table_idx}: No columns")
                continue

            # Check for inconsistent row structure
            expected_cells = len(table.rows[0].cells)
            for row_idx, row in enumerate(table.rows):
                if len(row.cells) != expected_cells:
                    table_issues.append(
                        f"Table {table_idx}, Row {row_idx}: Expected {expected_cells} cells, found {len(row.cells)}"
                    )

            # Check for completely empty tables
            has_content = False
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        has_content = True
                        break
                if has_content:
                    break

            if not has_content:
                table_issues.append(f"Table {table_idx}: Completely empty")

        if table_issues:
            self.fail(f"Table structure issues found:\n" + "\n".join(table_issues))

    def test_spacing_and_margins_reasonable(self):
        """Document should have reasonable spacing and margins."""
        # Check section margins
        for section in self.output_doc.sections:
            margin_issues = []

            # Check if margins are too small (less than 0.5 inches) or too large (more than 2 inches)
            if section.top_margin.inches < 0.3:
                margin_issues.append(
                    f"Top margin too small: {section.top_margin.inches:.2f} inches"
                )
            if section.top_margin.inches > 2.0:
                margin_issues.append(
                    f"Top margin too large: {section.top_margin.inches:.2f} inches"
                )

            if section.bottom_margin.inches < 0.3:
                margin_issues.append(
                    f"Bottom margin too small: {section.bottom_margin.inches:.2f} inches"
                )
            if section.bottom_margin.inches > 2.0:
                margin_issues.append(
                    f"Bottom margin too large: {section.bottom_margin.inches:.2f} inches"
                )

            if section.left_margin.inches < 0.3:
                margin_issues.append(
                    f"Left margin too small: {section.left_margin.inches:.2f} inches"
                )
            if section.left_margin.inches > 2.0:
                margin_issues.append(
                    f"Left margin too large: {section.left_margin.inches:.2f} inches"
                )

            if section.right_margin.inches < 0.3:
                margin_issues.append(
                    f"Right margin too small: {section.right_margin.inches:.2f} inches"
                )
            if section.right_margin.inches > 2.0:
                margin_issues.append(
                    f"Right margin too large: {section.right_margin.inches:.2f} inches"
                )

            if margin_issues:
                self.fail(f"Margin issues found:\n" + "\n".join(margin_issues))

    def test_text_formatting_quality(self):
        """Text should have appropriate formatting for a professional document."""
        formatting_issues = []

        # Check for mixed fonts within the document
        font_names = set()

        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.name:
                                font_names.add(run.font.name)

        for paragraph in self.output_doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name:
                    font_names.add(run.font.name)

        # Professional documents typically use 1-3 font families
        if len(font_names) > 5:
            formatting_issues.append(
                f"Too many different fonts used: {list(font_names)}"
            )

        # Check for reasonable text density (not too cramped)
        total_text_length = 0
        paragraph_count = 0

        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            total_text_length += len(paragraph.text)
                            paragraph_count += 1

        for paragraph in self.output_doc.paragraphs:
            if paragraph.text.strip():
                total_text_length += len(paragraph.text)
                paragraph_count += 1

        if paragraph_count > 0:
            avg_para_length = total_text_length / paragraph_count

            # Check for extremely long paragraphs (might indicate formatting issues)
            if avg_para_length > 500:
                formatting_issues.append(
                    f"Average paragraph length too long: {avg_para_length:.1f} characters"
                )

        if formatting_issues:
            self.fail(f"Text formatting issues found:\n" + "\n".join(formatting_issues))

    def test_professional_appearance_indicators(self):
        """Document should have indicators of professional formatting."""
        professional_indicators = {
            "has_bold_text": False,
            "has_varied_font_sizes": False,
            "has_proper_tables": False,
            "has_structured_content": False,
        }

        font_sizes = set()

        # Check for bold text and font size variety
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.bold:
                                professional_indicators["has_bold_text"] = True
                            if run.font.size:
                                font_sizes.add(run.font.size.pt)

        for paragraph in self.output_doc.paragraphs:
            for run in paragraph.runs:
                if run.bold:
                    professional_indicators["has_bold_text"] = True
                if run.font.size:
                    font_sizes.add(run.font.size.pt)

        if len(font_sizes) >= 2:
            professional_indicators["has_varied_font_sizes"] = True

        # Check for proper table structure
        if len(self.output_doc.tables) >= 3:
            professional_indicators["has_proper_tables"] = True

        # Check for structured content (headers, sections)
        content_keywords = [
            "experience",
            "education",
            "skills",
            "technical",
            "proficiencies",
        ]
        all_text = ""
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text.lower() + " "

        found_keywords = sum(1 for keyword in content_keywords if keyword in all_text)
        if found_keywords >= 3:
            professional_indicators["has_structured_content"] = True

        # Report on professional indicators
        missing_indicators = [
            key for key, value in professional_indicators.items() if not value
        ]

        if len(missing_indicators) > 2:
            self.fail(
                f"Document lacks professional formatting indicators. Missing: {missing_indicators}. "
                f"This may indicate poor visual formatting quality."
            )

    def test_compare_with_reference_structure(self):
        """Generated document should have similar structural complexity to reference."""
        output_stats = self._get_document_stats(self.output_doc)
        reference_stats = self._get_document_stats(self.reference_doc)

        # Compare key structural elements
        structure_issues = []

        # Table count should be similar (within 50% difference)
        table_diff = abs(output_stats["table_count"] - reference_stats["table_count"])
        if table_diff > reference_stats["table_count"] * 0.5:
            structure_issues.append(
                f"Table count very different: output={output_stats['table_count']}, "
                f"reference={reference_stats['table_count']}"
            )

        # Content length should be reasonably similar (within 30% difference)
        content_diff_ratio = (
            abs(
                output_stats["total_text_length"] - reference_stats["total_text_length"]
            )
            / reference_stats["total_text_length"]
        )
        if content_diff_ratio > 0.3:
            structure_issues.append(
                f"Content length very different: output={output_stats['total_text_length']}, "
                f"reference={reference_stats['total_text_length']} (diff: {content_diff_ratio:.1%})"
            )

        if structure_issues:
            self.fail(
                f"Structural differences from reference:\n"
                + "\n".join(structure_issues)
            )

    def _get_document_stats(self, doc):
        """Get statistical information about a document."""
        stats = {
            "table_count": len(doc.tables),
            "paragraph_count": len(doc.paragraphs),
            "total_text_length": 0,
            "font_sizes": set(),
            "has_bold": False,
        }

        # Count text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        stats["total_text_length"] += len(paragraph.text)
                        for run in paragraph.runs:
                            if run.bold:
                                stats["has_bold"] = True
                            if run.font.size:
                                stats["font_sizes"].add(run.font.size.pt)

        # Count text in paragraphs
        for paragraph in doc.paragraphs:
            stats["total_text_length"] += len(paragraph.text)
            for run in paragraph.runs:
                if run.bold:
                    stats["has_bold"] = True
                if run.font.size:
                    stats["font_sizes"].add(run.font.size.pt)

        return stats


if __name__ == "__main__":
    unittest.main()
