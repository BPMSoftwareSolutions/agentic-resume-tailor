"""Unit tests for export_docx.py - Markdown to DOCX conversion."""

import os
import sys
import tempfile
from pathlib import Path

# Add src to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

import unittest

from docx import Document

from export_docx import md_to_docx


class TestMdToDocx(unittest.TestCase):
    """Test the md_to_docx function."""

    def setUp(self):
        """Create temporary directory for test files."""
        self.temp_dir = tempfile.mkdtemp()

    def tearDown(self):
        """Clean up temporary files."""
        import shutil

        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)

    def test_creates_docx_file(self):
        """Should create a DOCX file."""
        md_path = Path(self.temp_dir) / "test.md"
        docx_path = Path(self.temp_dir) / "test.docx"

        md_path.write_text("# Test Resume\n\nSome content.", encoding="utf-8")
        md_to_docx(str(md_path), str(docx_path))

        self.assertTrue(docx_path.exists())

    def test_converts_h1_header(self):
        """Should skip H1 headers (name) per template format."""
        md_path = Path(self.temp_dir) / "test.md"
        docx_path = Path(self.temp_dir) / "test.docx"

        md_path.write_text("# John Doe\nDevOps Engineer", encoding="utf-8")
        md_to_docx(str(md_path), str(docx_path))

        doc = Document(str(docx_path))
        self.assertGreater(len(doc.paragraphs), 1)
        # First paragraph is empty (matching reference template)
        self.assertEqual("", doc.paragraphs[0].text.strip())
        # H1 is skipped, title line is shown in paragraph 1
        self.assertIn("DevOps Engineer", doc.paragraphs[1].text)

    def test_converts_h2_headers(self):
        """Should skip H2 headers (sections) per template format."""
        md_path = Path(self.temp_dir) / "test.md"
        docx_path = Path(self.temp_dir) / "test.docx"

        md_path.write_text(
            "Title\n## Summary\nSome summary text\n## Education\n- University Name",
            encoding="utf-8",
        )
        md_to_docx(str(md_path), str(docx_path))

        doc = Document(str(docx_path))
        text = "\n".join([p.text for p in doc.paragraphs])
        # H2 headers are skipped, but content is shown
        self.assertIn("Some summary text", text)
        self.assertIn("University Name", text)

    def test_converts_bullet_points(self):
        """Should convert experience bullet points to regular paragraphs."""
        md_path = Path(self.temp_dir) / "test.md"
        docx_path = Path(self.temp_dir) / "test.docx"

        md_content = "Title\n## Experience\n**Role — Company** (2020 – 2024)\n- Built Python automation\n- Managed CI/CD\n- Used Terraform"
        md_path.write_text(md_content, encoding="utf-8")
        md_to_docx(str(md_path), str(docx_path))

        doc = Document(str(docx_path))
        text = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("Python", text)
        self.assertIn("CI/CD", text)
        self.assertIn("Terraform", text)

    def test_converts_bold_text(self):
        """Should convert bold text (job titles)."""
        md_path = Path(self.temp_dir) / "test.md"
        docx_path = Path(self.temp_dir) / "test.docx"

        md_content = "**Senior Engineer — Company Name**"
        md_path.write_text(md_content, encoding="utf-8")
        md_to_docx(str(md_path), str(docx_path))

        doc = Document(str(docx_path))
        text = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("Senior Engineer", text)
        self.assertIn("Company Name", text)

    def test_preserves_special_characters(self):
        """Should preserve special characters like em-dash."""
        md_path = Path(self.temp_dir) / "test.md"
        docx_path = Path(self.temp_dir) / "test.docx"

        md_content = "Built pipeline — improved reliability"
        md_path.write_text(md_content, encoding="utf-8")
        md_to_docx(str(md_path), str(docx_path))

        doc = Document(str(docx_path))
        text = "\n".join([p.text for p in doc.paragraphs])
        self.assertIn("—", text)

    def test_handles_empty_lines(self):
        """Should handle empty lines gracefully."""
        md_path = Path(self.temp_dir) / "test.md"
        docx_path = Path(self.temp_dir) / "test.docx"

        md_content = "# Name\n\n\n## Section\n\nContent"
        md_path.write_text(md_content, encoding="utf-8")
        md_to_docx(str(md_path), str(docx_path))

        doc = Document(str(docx_path))
        self.assertGreater(len(doc.paragraphs), 0)

    def test_handles_utf8_encoding(self):
        """Should handle UTF-8 encoded characters."""
        md_path = Path(self.temp_dir) / "test.md"
        docx_path = Path(self.temp_dir) / "test.docx"

        md_content = "# José García\nSoftware Engineer\nC# and C++ developer"
        md_path.write_text(md_content, encoding="utf-8")
        md_to_docx(str(md_path), str(docx_path))

        doc = Document(str(docx_path))
        text = "\n".join([p.text for p in doc.paragraphs])
        # H1 is skipped, but title and content are shown
        self.assertIn("Software Engineer", text)
        self.assertIn("C#", text)

    def test_real_world_resume(self):
        """Test with a realistic resume structure."""
        md_path = Path(self.temp_dir) / "test.md"
        docx_path = Path(self.temp_dir) / "test.docx"

        md_content = """# John Doe
DevOps Engineer
Location • email@example.com • (123) 456-7890

## Summary
Experienced DevOps engineer with 10+ years.

## Skills
- CI/CD
- Python
- Terraform

## Experience
**Senior DevOps Engineer — Tech Company** (2020 – 2024)
- Built CI/CD pipelines improving deployment frequency 50%.
- Managed infrastructure with Terraform.

## Education
- University Name — Bachelor of Science (2010)
"""
        md_path.write_text(md_content, encoding="utf-8")
        md_to_docx(str(md_path), str(docx_path))

        doc = Document(str(docx_path))
        text = "\n".join([p.text for p in doc.paragraphs])

        # Check key content is present (H1 and H2 headers are skipped per template)
        self.assertIn("DevOps Engineer", text)
        self.assertIn("Location", text)
        self.assertIn("Experienced DevOps engineer", text)
        self.assertIn("Tech Company", text)
        self.assertIn("Senior DevOps Engineer", text)
        self.assertIn("50%", text)
        self.assertIn("University Name", text)

    def test_file_not_found_error(self):
        """Should raise error for non-existent markdown file."""
        docx_path = Path(self.temp_dir) / "test.docx"

        with self.assertRaises(FileNotFoundError):
            md_to_docx("nonexistent.md", str(docx_path))


if __name__ == "__main__":
    unittest.main()
