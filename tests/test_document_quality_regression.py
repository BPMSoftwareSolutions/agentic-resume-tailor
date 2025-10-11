"""
Comprehensive document quality regression tests.
These tests are designed to catch the type of visual quality issues that make
a document look like a 'hot mess' vs professional.
"""
import sys
from pathlib import Path
import tempfile
import os

# Add src to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / 'src'))

import unittest
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


class TestDocumentQualityRegression(unittest.TestCase):
    """Test that generated documents maintain professional quality standards."""
    
    def setUp(self):
        """Load the generated output and reference template."""
        # Use the generated baseline for testing
        self.output_path = Path(__file__).parent.parent / "out" / "test_generated_baseline.docx"
        # Use the original template as reference
        self.reference_path = Path(__file__).parent.parent / "out" / "Sidney Jones Resume - Solution Architect Leader.docx"

        if not self.output_path.exists():
            self.skipTest(f"Output file not found: {self.output_path}")

        if not self.reference_path.exists():
            self.skipTest(f"Reference file not found: {self.reference_path}")

        self.output_doc = Document(str(self.output_path))
        self.reference_doc = Document(str(self.reference_path))
    
    def test_prevents_cramped_layout_regression(self):
        """Prevent regression to cramped, unprofessional layout."""
        quality_issues = []
        
        # Check line spacing in tables (cramped layouts have minimal spacing)
        cramped_tables = 0
        for table_idx, table in enumerate(self.output_doc.tables):
            cell_spacing_issues = 0
            for row in table.rows:
                for cell in row.cells:
                    # Check if cell has reasonable content spacing
                    text_length = sum(len(p.text) for p in cell.paragraphs)
                    para_count = len([p for p in cell.paragraphs if p.text.strip()])
                    
                    if para_count > 0 and text_length > 100:  # Substantial content
                        avg_text_per_para = text_length / para_count
                        if avg_text_per_para > 200:  # Very long paragraphs might indicate cramped formatting
                            cell_spacing_issues += 1
            
            if cell_spacing_issues > len(table.rows) * 0.5:  # More than half the rows have issues
                cramped_tables += 1
        
        if cramped_tables > 2:
            quality_issues.append(f"Detected {cramped_tables} tables with cramped formatting")
        
        # Check for reasonable paragraph breaks
        total_paragraphs = len(self.output_doc.paragraphs)
        total_text = sum(len(p.text) for p in self.output_doc.paragraphs)
        
        if total_paragraphs > 0 and total_text > 0:
            avg_para_length = total_text / total_paragraphs
            if avg_para_length > 300:  # Very long average paragraphs
                quality_issues.append(f"Average paragraph length too long ({avg_para_length:.0f} chars) - may indicate cramped layout")
        
        if quality_issues:
            self.fail(f"Layout quality issues detected:\n" + "\n".join(quality_issues))
    
    def test_prevents_inconsistent_formatting_regression(self):
        """Prevent regression to inconsistent, messy formatting."""
        formatting_issues = []
        
        # Check for font consistency across the document
        fonts_used = set()
        font_sizes_used = set()
        
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.name:
                                fonts_used.add(run.font.name)
                            if run.font.size:
                                font_sizes_used.add(run.font.size.pt)
        
        for paragraph in self.output_doc.paragraphs:
            for run in paragraph.runs:
                if run.font.name:
                    fonts_used.add(run.font.name)
                if run.font.size:
                    font_sizes_used.add(run.font.size.pt)
        
        # Professional documents use limited font variety
        if len(fonts_used) > 4:
            formatting_issues.append(f"Too many fonts used ({len(fonts_used)}): {list(fonts_used)}")
        
        if len(font_sizes_used) > 6:
            formatting_issues.append(f"Too many font sizes used ({len(font_sizes_used)}): {sorted(font_sizes_used)}")
        
        # Check for wildly different font sizes (indicates poor hierarchy)
        if font_sizes_used:
            min_size = min(font_sizes_used)
            max_size = max(font_sizes_used)
            if max_size / min_size > 3:  # More than 3x difference
                formatting_issues.append(f"Extreme font size variation: {min_size}pt to {max_size}pt")
        
        if formatting_issues:
            self.fail(f"Formatting consistency issues detected:\n" + "\n".join(formatting_issues))
    
    def test_prevents_poor_visual_hierarchy_regression(self):
        """Prevent regression to poor visual hierarchy."""
        hierarchy_issues = []
        
        # Check that headers/titles are distinguished from body text
        has_size_variation = False
        has_bold_headers = False
        
        font_sizes = []
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.font.size:
                                font_sizes.append(run.font.size.pt)
                            if run.bold and run.font.size and run.font.size.pt > 10:
                                has_bold_headers = True
        
        if len(set(font_sizes)) >= 2:
            has_size_variation = True
        
        if not has_size_variation and not has_bold_headers:
            hierarchy_issues.append("No clear visual hierarchy detected (no size variation or bold headers)")
        
        # Check for proper sectioning
        all_text = ""
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += cell.text.lower() + " "
        
        # Should have clear sections
        expected_sections = ['technical', 'experience', 'education']
        found_sections = [section for section in expected_sections if section in all_text]
        
        if len(found_sections) < 2:
            hierarchy_issues.append(f"Poor content organization - only found sections: {found_sections}")
        
        if hierarchy_issues:
            self.fail(f"Visual hierarchy issues detected:\n" + "\n".join(hierarchy_issues))
    
    def test_prevents_unprofessional_spacing_regression(self):
        """Prevent regression to unprofessional spacing patterns."""
        spacing_issues = []
        
        # Check document margins
        for section in self.output_doc.sections:
            # Very small margins make documents look cramped
            if section.top_margin.inches < 0.4:
                spacing_issues.append(f"Top margin too small: {section.top_margin.inches:.2f} inches")
            if section.left_margin.inches < 0.4:
                spacing_issues.append(f"Left margin too small: {section.left_margin.inches:.2f} inches")
            
            # Very large margins waste space
            if section.top_margin.inches > 1.5:
                spacing_issues.append(f"Top margin too large: {section.top_margin.inches:.2f} inches")
            if section.left_margin.inches > 1.5:
                spacing_issues.append(f"Left margin too large: {section.left_margin.inches:.2f} inches")
        
        # Check for reasonable table spacing
        if len(self.output_doc.tables) > 0:
            # Tables should not be completely empty or have zero-width cells
            for table_idx, table in enumerate(self.output_doc.tables):
                empty_cells = 0
                total_cells = 0
                
                for row in table.rows:
                    for cell in row.cells:
                        total_cells += 1
                        if not cell.text.strip():
                            empty_cells += 1
                
                if total_cells > 0 and empty_cells / total_cells > 0.7:  # More than 70% empty
                    spacing_issues.append(f"Table {table_idx} is mostly empty ({empty_cells}/{total_cells} cells)")
        
        if spacing_issues:
            self.fail(f"Spacing quality issues detected:\n" + "\n".join(spacing_issues))
    
    def test_document_readability_standards(self):
        """Ensure document meets basic readability standards."""
        readability_issues = []
        
        # Extract all readable text
        all_text = ""
        word_count = 0
        
        for table in self.output_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        all_text += cell_text + " "
                        word_count += len(cell_text.split())
        
        for paragraph in self.output_doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                all_text += para_text + " "
                word_count += len(para_text.split())
        
        # Basic readability checks
        if word_count < 100:
            readability_issues.append(f"Document too short ({word_count} words) - may be incomplete")
        
        if word_count > 2000:
            readability_issues.append(f"Document very long ({word_count} words) - may need better formatting")
        
        # Check for reasonable sentence structure
        sentence_count = all_text.count('.') + all_text.count('!') + all_text.count('?')
        if sentence_count > 0:
            avg_words_per_sentence = word_count / sentence_count
            if avg_words_per_sentence > 40:
                readability_issues.append(f"Very long sentences (avg {avg_words_per_sentence:.1f} words) - hard to read")
        
        # Check for proper capitalization (indicates good formatting)
        if all_text:
            capital_ratio = sum(1 for c in all_text if c.isupper()) / len(all_text)
            if capital_ratio > 0.15:  # More than 15% capitals
                readability_issues.append(f"Excessive capitalization ({capital_ratio:.1%}) - may indicate formatting issues")
            if capital_ratio < 0.02:  # Less than 2% capitals
                readability_issues.append(f"Insufficient capitalization ({capital_ratio:.1%}) - may indicate formatting issues")
        
        if readability_issues:
            self.fail(f"Readability issues detected:\n" + "\n".join(readability_issues))
    
    def test_professional_quality_benchmarks(self):
        """Ensure document meets professional quality benchmarks."""
        quality_issues = []
        
        # Compare key metrics with reference document
        output_metrics = self._calculate_quality_metrics(self.output_doc)
        reference_metrics = self._calculate_quality_metrics(self.reference_doc)
        
        # Table structure should be similar
        table_diff = abs(output_metrics['table_count'] - reference_metrics['table_count'])
        if table_diff > 2:
            quality_issues.append(f"Table count significantly different: {output_metrics['table_count']} vs {reference_metrics['table_count']}")
        
        # Content density should be reasonable
        content_density_diff = abs(output_metrics['content_density'] - reference_metrics['content_density'])
        if content_density_diff > 50:  # More than 50 chars/para difference
            quality_issues.append(f"Content density very different: {output_metrics['content_density']:.1f} vs {reference_metrics['content_density']:.1f}")
        
        # Should have some formatting variation
        if not output_metrics['has_formatting_variety']:
            quality_issues.append("No formatting variety detected - document may look plain")
        
        # Should have reasonable structure complexity
        if output_metrics['structure_complexity'] < 3:
            quality_issues.append(f"Low structure complexity ({output_metrics['structure_complexity']}) - document may look simple")
        
        if quality_issues:
            self.fail(f"Professional quality issues detected:\n" + "\n".join(quality_issues))
    
    def _calculate_quality_metrics(self, doc):
        """Calculate quality metrics for a document."""
        metrics = {
            'table_count': len(doc.tables),
            'paragraph_count': 0,
            'total_chars': 0,
            'font_variety': set(),
            'size_variety': set(),
            'has_bold': False,
            'has_formatting_variety': False,
            'structure_complexity': 0
        }
        
        # Analyze content
        for table in doc.tables:
            metrics['structure_complexity'] += 1
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text.strip():
                            metrics['paragraph_count'] += 1
                            metrics['total_chars'] += len(paragraph.text)
                        
                        for run in paragraph.runs:
                            if run.font.name:
                                metrics['font_variety'].add(run.font.name)
                            if run.font.size:
                                metrics['size_variety'].add(run.font.size.pt)
                            if run.bold:
                                metrics['has_bold'] = True
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                metrics['paragraph_count'] += 1
                metrics['total_chars'] += len(paragraph.text)
            
            for run in paragraph.runs:
                if run.font.name:
                    metrics['font_variety'].add(run.font.name)
                if run.font.size:
                    metrics['size_variety'].add(run.font.size.pt)
                if run.bold:
                    metrics['has_bold'] = True
        
        # Calculate derived metrics
        if metrics['paragraph_count'] > 0:
            metrics['content_density'] = metrics['total_chars'] / metrics['paragraph_count']
        else:
            metrics['content_density'] = 0
        
        metrics['has_formatting_variety'] = (
            len(metrics['font_variety']) > 1 or
            len(metrics['size_variety']) > 1 or
            metrics['has_bold']
        )
        
        # Structure complexity: tables + font variety + size variety
        metrics['structure_complexity'] = (
            metrics['table_count'] +
            len(metrics['font_variety']) +
            len(metrics['size_variety'])
        )
        
        return metrics


if __name__ == '__main__':
    unittest.main()