"""
Content accuracy and formatting consistency tests.
These tests validate that the generated content matches the expected format
and catches issues like tab vs space differences, line formatting problems, etc.
"""

import os
import re
import sys
import tempfile
from pathlib import Path

# Add src to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

import difflib
import unittest

from docx import Document


class TestContentAccuracyValidation(unittest.TestCase):
    """Test that generated content has proper accuracy and formatting."""

    def setUp(self):
        """Load the generated output and reference template."""
        # Use the generated baseline for testing
        self.output_path = (
            Path(__file__).parent.parent / "out" / "test_generated_baseline.docx"
        )
        # Use the original template as reference
        self.reference_path = (
            Path(__file__).parent.parent
            / "out"
            / "Sidney Jones Resume - Solution Architect Leader.docx"
        )

        if not self.output_path.exists():
            self.skipTest(f"Output file not found: {self.output_path}")

        if not self.reference_path.exists():
            self.skipTest(f"Reference file not found: {self.reference_path}")

        self.output_doc = Document(str(self.output_path))
        self.reference_doc = Document(str(self.reference_path))

    def test_company_date_formatting_consistency(self):
        """Company names and dates should have consistent formatting."""
        all_text = self._extract_all_text(self.output_doc)

        # Look for company name and date patterns
        company_date_patterns = [
            r"BPM Software Solutions.*?20\d{2}\s*[–-]\s*20\d{2}",
            r"Interactive Business Solutions.*?20\d{2}\s*[–-]\s*20\d{2}",
            r"Soave Enterprises.*?20\d{2}\s*[–-]\s*20\d{2}",
            r"John Deere Landscapes.*?20\d{2}\s*[–-]\s*20\d{2}",
            r"CGI - Daugherty.*?20\d{2}\s*[–-]\s*20\d{2}",
        ]

        found_patterns = []
        formatting_issues = []

        for pattern in company_date_patterns:
            matches = re.finditer(pattern, all_text, re.IGNORECASE)
            for match in matches:
                found_text = match.group()
                found_patterns.append(found_text)

                # Check for inconsistent spacing (multiple spaces where one expected)
                if "      " in found_text:  # Multiple spaces
                    # This might be intentional alignment, but check context
                    if not self._is_intentional_alignment(found_text):
                        formatting_issues.append(
                            f"Excessive spacing in: '{found_text}'"
                        )

                # Check for tab characters that might not render properly
                if "\t" in found_text:
                    formatting_issues.append(f"Tab character found in: '{found_text}'")

        # Should find at least 4 company-date patterns
        self.assertGreaterEqual(
            len(found_patterns),
            4,
            f"Expected at least 4 company-date patterns, found {len(found_patterns)}: {found_patterns}",
        )

        if formatting_issues:
            self.fail(
                f"Company-date formatting issues:\n" + "\n".join(formatting_issues)
            )

    def test_job_title_section_structure(self):
        """Job titles and company sections should have proper structure."""
        all_text = self._extract_all_text(self.output_doc)
        lines = all_text.split("\n")

        structure_issues = []

        # Look for the CGI section specifically (as it was reformatted)
        cgi_section_found = False
        cgi_title_on_separate_line = False

        for i, line in enumerate(lines):
            line = line.strip()

            if "CGI - Daugherty" in line and "2022" in line:
                cgi_section_found = True

                # Check if the next line has the job title
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    if (
                        "Principle Consultant" in next_line
                        or "Principal Consultant" in next_line
                    ):
                        cgi_title_on_separate_line = True

        if cgi_section_found and not cgi_title_on_separate_line:
            structure_issues.append(
                "CGI section job title should be on separate line from company/dates"
            )

        # Check for other structural consistency
        company_lines = []
        for line in lines:
            line = line.strip()
            if any(
                company in line
                for company in [
                    "BPM Software Solutions",
                    "Interactive Business Solutions",
                    "Soave Enterprises",
                    "John Deere Landscapes",
                ]
            ):
                company_lines.append(line)

        # All company lines should have similar structure
        if len(company_lines) >= 2:
            first_structure = self._analyze_line_structure(company_lines[0])
            for i, line in enumerate(company_lines[1:], 1):
                line_structure = self._analyze_line_structure(line)
                if line_structure != first_structure:
                    structure_issues.append(
                        f"Inconsistent structure: '{company_lines[0]}' vs '{line}'"
                    )

        if structure_issues:
            self.fail(
                f"Job title section structure issues:\n" + "\n".join(structure_issues)
            )

    def test_certification_formatting_consistency(self):
        """Certifications should be properly formatted and complete."""
        all_text = self._extract_all_text(self.output_doc)

        # Expected certifications
        expected_certifications = [
            "SAFe 5 Certified DevOps Practitioner",
            "SAFe 5 Certified - Advanced Scrum Master",
            "GitHub Community Innovative Contributor",
            "Compuware President of ERG",
            "Detroit Community Youth Mentorship Award",
        ]

        missing_certifications = []
        formatting_issues = []

        for cert in expected_certifications:
            if cert not in all_text:
                # Check for partial matches (might be formatting issue)
                cert_words = cert.split()
                if len(cert_words) > 1:
                    partial_match = any(word in all_text for word in cert_words[:2])
                    if partial_match:
                        formatting_issues.append(
                            f"Partial match for certification: '{cert}'"
                        )
                    else:
                        missing_certifications.append(cert)
                else:
                    missing_certifications.append(cert)

        # Check if certifications are on separate lines (better formatting)
        lines = all_text.split("\n")
        combined_cert_lines = []

        for line in lines:
            line = line.strip()
            # Look for lines with multiple certifications combined
            cert_count = sum(1 for cert in expected_certifications if cert in line)
            if cert_count > 1:
                combined_cert_lines.append(line)

        if combined_cert_lines:
            formatting_issues.append(
                f"Certifications combined on single lines: {combined_cert_lines}"
            )

        issues = []
        if missing_certifications:
            issues.append(f"Missing certifications: {missing_certifications}")
        if formatting_issues:
            issues.extend(formatting_issues)

        if issues:
            self.fail(f"Certification formatting issues:\n" + "\n".join(issues))

    def test_content_completeness_vs_reference(self):
        """Generated content should be complete compared to reference."""
        output_text = self._extract_all_text(self.output_doc)
        reference_text = self._extract_all_text(self.reference_doc)

        # Split into lines for comparison
        output_lines = [
            line.strip() for line in output_text.split("\n") if line.strip()
        ]
        reference_lines = [
            line.strip() for line in reference_text.split("\n") if line.strip()
        ]

        # Check for significant content differences
        content_issues = []

        # Character count difference
        char_diff = abs(len(output_text) - len(reference_text))
        char_diff_percent = char_diff / len(reference_text) * 100

        if char_diff_percent > 5:  # More than 5% difference
            content_issues.append(
                f"Significant content length difference: {char_diff} characters ({char_diff_percent:.1f}%)"
            )

        # Line count difference
        line_diff = abs(len(output_lines) - len(reference_lines))
        line_diff_percent = line_diff / len(reference_lines) * 100

        if line_diff_percent > 10:  # More than 10% difference in lines
            content_issues.append(
                f"Significant line count difference: {line_diff} lines ({line_diff_percent:.1f}%)"
            )

        # Check for missing key content sections
        key_sections = [
            "Software Engineering & Solutions Architect Leader",
            "Technical Proficiencies",
            "Areas of Expertise",
            "Career Experience",
            "Education",
        ]

        missing_sections = []
        for section in key_sections:
            if section not in output_text:
                missing_sections.append(section)

        if missing_sections:
            content_issues.append(f"Missing key sections: {missing_sections}")

        if content_issues:
            self.fail(f"Content completeness issues:\n" + "\n".join(content_issues))

    def test_whitespace_and_spacing_consistency(self):
        """Whitespace and spacing should be consistent throughout document."""
        all_text = self._extract_all_text(self.output_doc)
        lines = all_text.split("\n")

        spacing_issues = []

        # Check for inconsistent spacing patterns
        for i, line in enumerate(lines):
            if line.strip():  # Skip empty lines
                # Check for excessive spaces (more than 6 consecutive spaces)
                if "       " in line:  # 7+ spaces
                    spacing_issues.append(
                        f"Line {i+1}: Excessive spacing - '{line[:50]}...'"
                    )

                # Check for mixed tab/space patterns
                if "\t" in line and "  " in line:
                    spacing_issues.append(
                        f"Line {i+1}: Mixed tabs and spaces - '{line[:50]}...'"
                    )

                # Check for trailing whitespace (might indicate formatting issues)
                if line != line.rstrip():
                    spacing_issues.append(f"Line {i+1}: Trailing whitespace")

        # Limit reported issues to avoid overwhelming output
        if len(spacing_issues) > 10:
            spacing_issues = spacing_issues[:10] + [
                f"... and {len(spacing_issues) - 10} more spacing issues"
            ]

        if spacing_issues:
            self.fail(f"Whitespace consistency issues:\n" + "\n".join(spacing_issues))

    def _extract_all_text(self, doc):
        """Extract all text from a document."""
        all_text = ""

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        all_text += paragraph.text + "\n"

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            all_text += paragraph.text + "\n"

        return all_text

    def _is_intentional_alignment(self, text):
        """Check if spacing appears to be intentional alignment."""
        # If it follows a pattern like "Company Name      Date", it might be intentional
        return bool(re.search(r"\w+\s{4,}\d{4}", text))

    def _analyze_line_structure(self, line):
        """Analyze the structure of a line (for consistency checking)."""
        structure = {
            "has_company": bool(re.search(r"[A-Z][a-z]+ [A-Z][a-z]+", line)),
            "has_dates": bool(re.search(r"20\d{2}\s*[–-]\s*20\d{2}", line)),
            "has_location": bool(re.search(r", [A-Z]{2}", line)),
            "spacing_pattern": (
                "tabs" if "\t" in line else "spaces" if "    " in line else "minimal"
            ),
        }
        return structure


if __name__ == "__main__":
    unittest.main()
